from __future__ import annotations

from pathlib import Path
from typing import Optional

import pandas as pd


def require_xlsxwriter() -> None:
    try:
        import xlsxwriter  # noqa: F401
    except Exception as e:  # pragma: no cover
        raise RuntimeError("xlsxwriter is required to write XLSX reports") from e


def require_docx() -> None:
    try:
        import docx  # noqa: F401
    except Exception as e:  # pragma: no cover
        raise RuntimeError("python-docx is required to write DOCX reports") from e


def write_xlsx_table(
    *,
    df: pd.DataFrame,
    out_xlsx: Path,
    sample: str,
    org: str = "None",
    genomes: str = "None",
) -> None:
    require_xlsxwriter()
    import xlsxwriter

    out_xlsx.parent.mkdir(parents=True, exist_ok=True)

    cols = [
        "iter",
        "iter_cluster",
        "rank",
        "orig",
        "hit_count",
        "size, %",
        "pdiv_mean",
        "pdiv_median",
        "pdiv_q05",
        "pdiv_q95",
        "best_hit",
        "best_species",
        "best_name",
        "best_evalue",
        "best_pident",
        "coverage",
        "coverage_type",
        "task_used",
        "TAREAN_annotation",
        "cons_len",
        "seq",
        "pic_path",
    ]

    extra_hit_cols = sorted([c for c in df.columns if c.startswith("hit_count_")])
    extra_size_cols = sorted([c for c in df.columns if c.startswith("size, %_")])
    for c in extra_hit_cols:
        if c not in cols:
            idx = cols.index("hit_count") + 1
            cols.insert(idx, c)
    for c in extra_size_cols:
        if c not in cols:
            idx = cols.index("size, %") + 1
            cols.insert(idx, c)

    for c in cols:
        if c not in df.columns:
            df[c] = pd.NA
    df2 = df[cols].copy()

    wb = xlsxwriter.Workbook(str(out_xlsx))
    ws = wb.add_worksheet("satminer")

    header_fmt = wb.add_format({"bold": True, "bg_color": "#E6E6E6", "border": 1})
    mono_fmt = wb.add_format({"font_name": "Courier New", "font_size": 9})
    num2_fmt = wb.add_format({"num_format": "0.00"})

    ws.write(0, 0, "Sample", header_fmt)
    ws.write(0, 1, str(sample), header_fmt)
    ws.write(0, 2, "ORG", header_fmt)
    ws.write(0, 3, str(org) or "None", header_fmt)
    ws.write(0, 4, "GENOMES", header_fmt)
    ws.write(0, 5, str(genomes) or "None", header_fmt)

    start_row = 2
    for j, c in enumerate(cols):
        ws.write(start_row, j, c, header_fmt)

    ws.set_column(0, 0, 6)  # iter
    ws.set_column(1, 1, 14)  # iter_cluster
    ws.set_column(2, 2, 6)  # rank
    ws.set_column(3, 3, 28)  # orig
    ws.set_column(4, 6, 10)  # hit_count, size,%
    ws.set_column(7, 10, 12)  # pdiv stats
    ws.set_column(11, 11, 30)  # best_hit
    ws.set_column(12, 12, 18)  # best_species
    ws.set_column(13, 13, 40)  # best_name
    ws.set_column(14, 17, 12)  # blast metrics
    ws.set_column(18, 18, 28)  # annotation
    ws.set_column(19, 19, 9)  # cons_len
    ws.set_column(20, 20, 80)  # seq
    ws.set_column(21, 21, 40)  # pic_path (for debugging)

    r0 = start_row + 1
    img_col = cols.index("pic_path")
    for i, row in enumerate(df2.itertuples(index=False), start=0):
        r = r0 + i
        ws.set_row(r, 80)
        for j, val in enumerate(row):
            if j == img_col:
                p = "" if pd.isna(val) else str(val)
                ws.write(r, j, p, mono_fmt)
                if p and Path(p).exists():
                    try:
                        ws.insert_image(r, j, p, {"x_scale": 0.20, "y_scale": 0.20, "object_position": 1})
                    except Exception:
                        pass
                continue

            if pd.isna(val):
                continue
            if cols[j] in ("size, %", "pdiv_mean", "pdiv_median", "pdiv_q05", "pdiv_q95", "coverage", "best_pident"):
                try:
                    ws.write_number(r, j, float(val), num2_fmt)
                    continue
                except Exception:
                    pass
            ws.write(r, j, str(val))

    wb.close()


def write_docx_table(
    *,
    df: pd.DataFrame,
    out_docx: Path,
    sample: str,
    org: str = "None",
    genomes: str = "None",
    embed_top_images: int = 10,
) -> None:
    require_docx()
    import docx

    out_docx.parent.mkdir(parents=True, exist_ok=True)

    doc = docx.Document()
    doc.add_heading(f"satMiner report: {sample}", level=0)
    doc.add_paragraph(f"ORG: {org or 'None'}")
    doc.add_paragraph(f"GENOMES: {genomes or 'None'}")

    doc.add_heading("Summary (top by abundance)", level=1)
    if "size, %" in df.columns:
        df["size, %"] = pd.to_numeric(df["size, %"], errors="coerce")
    top = df.sort_values(["size, %"], ascending=False, na_position="last").head(20).copy()

    table = doc.add_table(rows=1, cols=8)
    hdr = table.rows[0].cells
    headers = ["iter", "cluster", "rank", "size,%", "hit_count", "best_hit", "coverage", "coverage_type"]
    for i, h in enumerate(headers):
        hdr[i].text = h
    for _, r in top.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(r.get("iter", ""))
        cells[1].text = str(r.get("iter_cluster", r.get("cluster", "")))
        cells[2].text = str(r.get("rank", ""))
        cells[3].text = str(r.get("size, %", ""))
        cells[4].text = str(r.get("hit_count", ""))
        cells[5].text = str(r.get("best_hit", ""))
        cells[6].text = str(r.get("coverage", ""))
        cells[7].text = str(r.get("coverage_type", ""))

        extra = []
        for c in sorted([c for c in df.columns if c.startswith("size, %_")]):
            extra.append(f"{c}={r.get(c,'')}")
        if extra:
            doc.add_paragraph("Per-sample: " + "; ".join(extra))

    doc.add_heading("By iteration", level=1)
    iters = sorted([int(x) for x in pd.to_numeric(df.get("iter"), errors="coerce").dropna().unique().tolist()] or [])
    for it in iters:
        sub = df[df["iter"] == it].copy()
        doc.add_heading(f"ITER={it}", level=2)
        sub = sub.sort_values(["size, %"], ascending=False, na_position="last").head(10)
        for _, r in sub.iterrows():
            doc.add_paragraph(
                f"{r.get('iter_cluster','')}: size,%={r.get('size, %','')} hit_count={r.get('hit_count','')} "
                f"best_hit={r.get('best_hit','')} cov={r.get('coverage','')} ({r.get('coverage_type','')})"
            )

    nimg = max(0, int(embed_top_images))
    if nimg:
        doc.add_heading(f"Top {nimg} cluster graphs", level=1)
        for _, r in top.head(nimg).iterrows():
            p = str(r.get("pic_path", "") or "")
            if not p:
                continue
            img = Path(p)
            if not img.exists():
                continue
            doc.add_paragraph(str(r.get("iter_cluster", "")))
            try:
                doc.add_picture(str(img), width=docx.shared.Inches(5.5))
            except Exception:
                doc.add_paragraph(f"(image path) {img}")

    doc.save(str(out_docx))


def legacy_render_report_via_repeatanalyzer(
    *,
    project: str,
    sample: str,
    df: pd.DataFrame,
    out_xlsx: Path,
    out_docx: Path,
) -> None:
    # Used by satminer_render_xlsx_docx_report.py to preserve historical behavior.
    import sys

    repo_root = Path(__file__).resolve().parents[2]
    if str(repo_root) not in sys.path:
        sys.path.insert(0, str(repo_root))

    from post_tarean.pipeline import RepeatAnalyzer

    analyzer = RepeatAnalyzer(project_id=project)
    ok_xlsx = analyzer.create_excel_report(df, str(out_xlsx), tarean_path=None)
    if not ok_xlsx:
        raise RuntimeError("Failed to create XLSX report (xlsxwriter missing?)")

    ok_docx = analyzer.create_word_report(sample, df, str(out_docx))
    if not ok_docx:
        raise RuntimeError("Failed to create DOCX report (python-docx missing?)")

