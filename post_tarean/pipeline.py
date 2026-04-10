#!/usr/bin/env python3
"""
RepOrtR Post-TAREAN Pipeline
Consolidated pipeline for repeat analysis after TAREAN processing

HOLY WORKFLOW COMPLIANCE: This module supports the modular workflow's analysis rules.
All configuration comes from projects/global_config.yaml.
"""

import os
import glob
import re
import argparse
import pandas as pd
import numpy as np
import textwrap
import tempfile
import subprocess
import traceback
from collections import defaultdict
from pathlib import Path
import logging
import sys

from post_tarean.config import ConfigManager, create_config_for_project
from post_tarean.io_helpers import resolve_tarean_path
from post_tarean.orchestrator import PostTareanOrchestrator, build_default_steps
from post_tarean.steps.base import PostTareanContext

# Optional imports
try:
    import xlsxwriter
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import plotly.express as px
    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False

try:
    from rapidfuzz import fuzz
    FUZZ_AVAILABLE = True
except ImportError:
    FUZZ_AVAILABLE = False

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class RepeatAnalyzer:
    """Main class for repeat analysis pipeline"""
    
    def __init__(self, config=None, project_id=None):
        """
        Initialise the analyzer.

        - If a config object (PipelineConfig) is provided, wrap it in a
          ConfigManager instance and ignore project_id (the caller is
          responsible for choosing the right project-specific config).
        - Otherwise, if project_id is given, load config from
          projects/global_config.yaml via create_config_for_project.
        - As a fallback, use the default ConfigManager() configuration.
        """
        if config is not None:
            # Wrap existing PipelineConfig in a fresh manager, but preserve the
            # calling project_id so that downstream helpers (like
            # resolve_tarean_path) can locate project-scoped data such as
            # TAREAN outputs and NCBI-derived repeat databases.
            mgr = ConfigManager(project_id=project_id)
            mgr.config = config
            # Attach project_id back onto the config object for convenience.
            if project_id is not None and not hasattr(mgr.config, "project_id"):
                setattr(mgr.config, "project_id", project_id)
            self.config_manager = mgr
        elif project_id:
            self.config_manager = create_config_for_project(project_id)
        else:
            self.config_manager = ConfigManager()
        
        self.config = self.config_manager.config
        self.depth_dict = self._get_depth_dict()
        
        # Get misc categories from config
        self.misc_categories = self.config.filtering.misc_categories
        
    def _get_depth_dict(self):
        """Get the depth dictionary for repeat classification"""
        return {
            ("Unclassified_repeat (conflicting evidences)", 0),
            ("|--rDNA", 1),
            ("|   |--45S_rDNA", 2),
            ("|   |   |--18S_rDNA", 3),
            ("|   |   |--25S_rDNA", 3),
            ("|   |   '--5.8S_rDNA", 3),
            ("|   '--5S_rDNA", 2),
            ("|--satellite", 1),
            ("'--mobile_element", 1),
            ("|--Class_I", 2),
            ("|   |--SINE", 3),
            ("|   |--LTR", 3),
            ("|   |   |--Ty1_copia", 4),
            ("|   |   |   |--Ale", 5),
            ("|   |   |   |--Alesia", 5),
            ("|   |   |   |--Angela", 5),
            ("|   |   |   |--Bianca", 5),
            ("|   |   |   |--Bryco", 5),
            ("|   |   |   |--Lyco", 5),
            ("|   |   |   |--Gymco-III", 5),
            ("|   |   |   |--Gymco-I", 5),
            ("|   |   |   |--Gymco-II", 5),
            ("|   |   |   |--Ikeros", 5),
            ("|   |   |   |--Ivana", 5),
            ("|   |   |   |--Gymco-IV", 5),
            ("|   |   |   |--Osser", 5),
            ("|   |   |   |--SIRE", 5),
            ("|   |   |   |--TAR", 5),
            ("|   |   |   |--Tork", 5),
            ("|   |   |   '--Ty1-outgroup", 5),
            ("|   |   '--Ty3_gypsy", 4),
            ("|   |       |--non-chromovirus", 5),
            ("|   |       |   |--non-chromo-outgroup", 6),
            ("|   |       |   |--Phygy", 6),
            ("|   |       |   |--Selgy", 6),
            ("|   |       |   '--OTA", 6),
            ("|   |       |       |--Athila", 7),
            ("|   |       |       '--Tat", 7),
            ("|   |       |           |--TatI", 8),
            ("|   |       |           |--TatII", 8),
            ("|   |       |           |--TatIII", 8),
            ("|   |       |           |--Ogre", 8),
            ("|   |       |           '--Retand", 8),
            ("|   |       '--chromovirus", 5),
            ("|   |           |--Chlamyvir", 6),
            ("|   |           |--Tcn1", 6),
            ("|   |           |--chromo-outgroup", 6),
            ("|   |           |--CRM", 6),
            ("|   |           |--Galadriel", 6),
            ("|   |           |--Tekay", 6),
            ("|   |           |--Reina", 6),
            ("|   |           '--chromo-unclass", 6),
            ("|   |--pararetrovirus", 3),
            ("|   |--DIRS", 3),
            ("|   |--Penelope", 3),
            ("|   '--LINE", 3),
            ("'--Class_II", 2),
            ("|--Subclass_1", 3),
            ("|   '--TIR", 4),
            ("|       |--MITE", 5),
            ("|       |--EnSpm_CACTA", 5),
            ("|       |--hAT", 5),
            ("|       |--Kolobok", 5),
            ("|       |--Merlin", 5),
            ("|       |--MuDR_Mutator", 5),
            ("|       |--Novosib", 5),
            ("|       |--P", 5),
            ("|       |--PIF_Harbinger", 5),
            ("|       |--PiggyBac", 5),
            ("|       |--Sola1", 5),
            ("|       |--Sola2", 5),
            ("|       '--Tc1_Mariner", 5),
            ("'--Subclass_2", 3),
            ("'--Helitron", 4),
            ("organelle", 0),
            ("|--plastid", 1),
            ("'--mitochondria", 1),
            ("Unclassified repeat (No evidence)", 0),
            ("contamination", 0),
        }
    
    def get_cypher_mapping(self):
        """Legacy metadata mapping (config-driven; cypher.csv is forbidden)."""
        from post_tarean.utils import get_cyphers
        return get_cyphers()
    
    def parse_tarean_data(self, path, index):
        """Parse TAREAN output data"""
        from post_tarean.utils import get_tareans
        return get_tareans(path, index)
    
    def parse_copy_data(self, path):
        """Parse copy number data from CLUSTER_TABLE.csv"""
        from post_tarean.utils import parse_copy_data as utils_parse_copy_data
        return utils_parse_copy_data(path)
    
    def parse_comparative_data(self, index, path):
        """Parse comparative analysis data"""
        from post_tarean.utils import parse_comparative_data as utils_parse_comparative_data

        return utils_parse_comparative_data(index, path)
    
    def run_blast_analysis(self, subjects, dbs=None, tasks=None, num_threads=None):
        """Run BLAST analysis using the consolidated blast module"""
        from post_tarean.blast_consolidated import BLASTAnalyzer
        
        # Initialise analyzer with the current ConfigManager so that
        # project-specific database paths (including NCBI-derived repeat
        # databases for Triticeae) are honoured, and so that the post-TAREAN
        # query-construction path (using TAREAN consensus FASTAs as queries
        # against external BLAST DBs) is enabled.
        blast_analyzer = BLASTAnalyzer(
            config=self.config,
            config_manager=self.config_manager,
            project_id=getattr(self.config_manager, "project_id", None),
        )
        
        # Use config values if not provided
        if dbs is None:
            dbs = self.config.blast.default_dbs
        if tasks is None:
            tasks = self.config.blast.default_tasks
        if num_threads is None:
            num_threads = self.config.blast.default_threads
        return blast_analyzer.run_blast(
            subjects=subjects,
            dbs=dbs,
            tasks=tasks,
            num_threads=num_threads
        )
    
    def generate_summary_report(self, index, path, blast_results=None):
        """Generate comprehensive summary report"""
        logger.info(f"Generating summary report for {index}")
        
        # Parse data (strict: failures must propagate so Snakemake can fail early).
        logger.info("Parsing TAREAN data...")
        tareans_data = self.parse_tarean_data(path, index)
        logger.info(f"TAREAN data shape: {tareans_data.shape}, columns: {list(tareans_data.columns)}")

        logger.info("Parsing copy data...")
        copy_data = self.parse_copy_data(path)
        logger.info(f"Copy data shape: {copy_data.shape}, columns: {list(copy_data.columns)}")

        logger.info("Parsing comparative data...")
        comparative_data = self.parse_comparative_data(index, path)
        logger.info(f"Comparative data shape: {comparative_data.shape}, columns: {list(comparative_data.columns)}")

        # Merge copy data (required for size/abundance columns in legacy reports).
        if "Cluster" not in copy_data.columns:
            raise KeyError("CLUSTER_TABLE.csv is missing required column 'Cluster'")
        logger.info("Merging TAREAN and copy data...")
        merged_data = pd.merge(tareans_data, copy_data, on="Cluster", how="left")
        logger.info(f"Merged data shape: {merged_data.shape}, columns: {list(merged_data.columns)}")

        # Merge comparative abundance/percentage columns (if comparative table is present).
        if (
            not comparative_data.empty
            and "Cluster" in comparative_data.columns
            and "Cluster" in merged_data.columns
        ):
            logger.info("Merging comparative data...")
            merged_data = pd.merge(merged_data, comparative_data, on="Cluster", how="left")

        # Add BLAST results if available
        if blast_results is not None and not blast_results.empty:
            logger.info("Adding BLAST results...")
            blast_summary = self.summarize_blast_results(blast_results, index)
            if not blast_summary.empty:
                merged_data = pd.merge(merged_data, blast_summary, on="Cluster", how="left")

        if merged_data.empty:
            raise RuntimeError("generate_summary_report produced empty merged table (unexpected)")

        if "size, %" not in merged_data.columns:
            raise KeyError("Merged report table is missing required column 'size, %'")
        if "TAREAN_annotation" not in merged_data.columns:
            raise KeyError("Merged report table is missing required column 'TAREAN_annotation'")

        return merged_data
    
    def summarize_blast_results(self, blast_results, index):
        """Summarize BLAST results for a given index with multiple matches and coverage analysis"""
        if blast_results.empty:
            return pd.DataFrame()
        
        # Debug: show what sequences are in the BLAST results
        logger.info(f"BLAST results contain {len(blast_results)} total hits")
        logger.info(f"Unique query sequences: {blast_results['qseqid'].unique()[:10]}")
        
        # Filter for the specific index - look for sequences that contain the index
        index_results = blast_results[blast_results['qseqid'].str.contains(index)]
        
        if index_results.empty:
            logger.warning(f"No BLAST results found for index {index}")
            return pd.DataFrame()
        
        # Group by cluster and analyze all matches
        summary = []
        for cluster, group in index_results.groupby('qseqid'):
            # Sort by E-value and get multiple significant matches
            significant_matches = group[group['evalue'] <= 0.001].sort_values('evalue')
            
            if len(significant_matches) > 0:
                # Get the best match
                best_hit = significant_matches.iloc[0]
                
                # Calculate coverage
                coverage = best_hit['length'] / best_hit['qlength']
                
                # Classify coverage type
                if coverage >= 0.95:
                    coverage_type = 'near-full'
                elif coverage >= 0.8:
                    coverage_type = 'composite'
                elif coverage >= 0.6:
                    coverage_type = 'partial'
                else:
                    coverage_type = 'weak'
                
                # Get task description
                task_desc = {
                    'megablast': 'high homology',
                    'dc-megablast': 'moderate homology',
                    'blastn': 'general homology'
                }.get(best_hit['task'], 'homology')
                
                summary.append({
                    'Cluster': cluster,  # Use uppercase to match merged data
                    'best_hit': best_hit['sseqid'],
                    'best_evalue': best_hit['evalue'],
                    'best_pident': best_hit['pident'],
                    'coverage': coverage,
                    'coverage_type': coverage_type,
                    'task_type': task_desc,
                    'num_hits': len(group),
                    'num_significant': len(significant_matches)
                })
            else:
                # No significant matches
                summary.append({
                    'Cluster': cluster,
                    'best_hit': 'No significant hits',
                    'best_evalue': float('inf'),
                    'best_pident': 0.0,
                    'coverage': 0.0,
                    'coverage_type': 'none',
                    'task_type': 'none',
                    'num_hits': len(group),
                    'num_significant': 0
                })
        
        logger.info(f"Found {len(summary)} BLAST hits for {index}")
        return pd.DataFrame(summary)
    
    def create_excel_report(self, data, output_file, tarean_path=None):
        """Create Excel report matching the original KP2.xlsx structure"""
        if not XLSX_AVAILABLE:
            logger.error("xlsxwriter not available, cannot create Excel report")
            return False
        
        # If TAREAN_annotation is missing (e.g. very minimal or legacy inputs),
        # inject a placeholder category so that the grouping logic below does
        # not fail with KeyError.
        safe_data = data.copy()
        if "TAREAN_annotation" not in safe_data.columns:
            safe_data["TAREAN_annotation"] = "Unclassified repeat (No evidence)"

        # Replace NaN/Inf with string for all values
        safe_data = safe_data.applymap(
            lambda x: str(x)
            if pd.isna(x)
            or (isinstance(x, float) and (pd.isna(x) or pd.isnull(x) or x == float("inf") or x == float("-inf")))
            else x
        )
        
        with xlsxwriter.Workbook(output_file) as workbook:
            worksheet = workbook.add_worksheet('main_sheet')
            
            # Define formats
            merge_format = workbook.add_format({
                'bold': True,
                'border': 1,
                'align': 'center',
                'valign': 'vcenter',
                'bg_color': '#E6E6FA'
            })
            
            merge_format_seq = workbook.add_format({
                'bold': True,
                'border': 1,
                'align': 'left',
                'valign': 'vcenter',
                'bg_color': '#F0F8FF'
            })
            
            header_format = workbook.add_format({
                'bold': True,
                'border': 1,
                'align': 'center',
                'valign': 'vcenter',
                'bg_color': '#D3D3D3'
            })
            
            # Side columns for metadata (2 columns)
            sidecols = 2
            
            # Define column structure based on available data
            annotation_columns = ['cluster', 'seq', 'pic_path', 'size, %', 'TAREAN_annotation']
            blast_columns = ['best_hit', 'best_evalue', 'best_pident', 'coverage', 'coverage_type', 'task_type']
            
            # Calculate column positions
            coords = {
                'annotation': sidecols,
                'blast': sidecols + len(annotation_columns)
            }
            
            # Write headers
            all_columns = annotation_columns + blast_columns
            for num, col_name in enumerate(all_columns):
                worksheet.write(0, num + sidecols, col_name, header_format)
            
            # Group data by TAREAN annotation type
            type_order = ['Putative satellites (high confidence)', 'Putative satellites (low confidence)', 'Putative LTR elements']
            
            row = 1
            for annotation_type in type_order:
                # Filter data for this annotation type
                type_data = safe_data[safe_data['TAREAN_annotation'] == annotation_type]
                
                if len(type_data) == 0:
                    continue
                
                # Merge range for annotation type header
                worksheet.merge_range(row, min(coords.values()), row, max(coords.values()) + len(all_columns) - 1, 
                                   annotation_type, merge_format)
                row += 1
                
                # Process each cluster in this type
                for _, cluster_row in type_data.iterrows():
                    # Handle both 'cluster' and 'Cluster' column names
                    cluster_col = 'cluster' if 'cluster' in cluster_row else 'Cluster'
                    cluster_name = cluster_row[cluster_col]
                    
                    # Create annotation chunk with correct column mapping
                    annotation_data = {}
                    for col in annotation_columns:
                        if col in cluster_row:
                            annotation_data[col] = cluster_row[col]
                        else:
                            annotation_data[col] = ''
                    annotation_chunk = pd.DataFrame([annotation_data])
                    
                    # Create BLAST chunk if available
                    blast_chunk = pd.DataFrame()
                    if all(col in cluster_row for col in blast_columns):
                        blast_data = {}
                        for col in blast_columns:
                            blast_data[col] = cluster_row[col]
                        blast_chunk = pd.DataFrame([blast_data])
                    
                    # Determine max rows for merging
                    max_rows = max(len(annotation_chunk), len(blast_chunk)) if len(blast_chunk) > 0 else len(annotation_chunk)
                    
                    # Write annotation data
                    col = coords['annotation']
                    for rel_row, data in annotation_chunk.iterrows():
                        for rel_col, item in enumerate(data):
                            if pd.isna(item) or item == '':
                                continue
                            
                            # Handle image path specially
                            if annotation_columns[rel_col] == 'pic_path':
                                path = item
                                item = ''
                            
                            # Merge cells for sequence data
                            if annotation_columns[rel_col] == 'seq':
                                if max_rows > 1:
                                    worksheet.merge_range(row + rel_row, col + rel_col, 
                                                       row + max_rows - 1, col + rel_col, 
                                                       item, merge_format_seq)
                                else:
                                    worksheet.write(row + rel_row, col + rel_col, item, merge_format_seq)
                                
                                # Insert image if available
                                if 'pic_path' in cluster_row and cluster_row['pic_path'] and os.path.exists(cluster_row['pic_path']):
                                    try:
                                        worksheet.insert_image(row + rel_row, col + rel_col, 
                                                            cluster_row['pic_path'], 
                                                            {'x_scale': 0.15, 'y_scale': 0.15, 'object_position': 1})
                                    except Exception as e:
                                        logger.warning(f"Could not insert image {cluster_row['pic_path']}: {e}")
                            else:
                                worksheet.write(row + rel_row, col + rel_col, item)
                    
                    # Write BLAST data if available
                    if len(blast_chunk) > 0:
                        col = coords['blast']
                        for rel_row, data in blast_chunk.iterrows():
                            for rel_col, item in enumerate(data):
                                if pd.isna(item) or item == '':
                                    continue
                                worksheet.write(row + rel_row, col + rel_col, item)
                    
                    row += max_rows
            
            # Add metadata in side columns
            worksheet.write(1, 0, 'Парных ридов подано на RepEx')
            worksheet.write(2, 0, 'Процент поданных ридов в топе кластеров')
            
            # Try to read metadata from index.html if a TAREAN path is available.
            try:
                index_html_path = None
                if tarean_path:
                    candidate = Path(str(tarean_path)) / "index.html"
                    if candidate.exists():
                        index_html_path = str(candidate)
                    else:
                        # Fallback to any index.html under the TAREAN directory.
                        hits = list(Path(str(tarean_path)).glob("**/index.html"))
                        if hits:
                            index_html_path = str(hits[0])

                if not index_html_path:
                    raise FileNotFoundError("index.html not found")

                with open(index_html_path, 'r') as file:
                    text = file.read()
                    pair_number_match = re.findall(r"<p class='character'>Number of analyzed reads: ([0-9\.]+)</p>", text)
                    percent_match = re.findall(r"<p class='character'>Proportion of reads in top clusters : ([0-9\.]+) %</p>", text)
                    
                    if pair_number_match:
                        worksheet.write(1, 1, int(pair_number_match[0]))
                    if percent_match:
                        worksheet.write(2, 1, int(percent_match[0]))
            except Exception as e:
                logger.warning(f"Could not read metadata from index.html: {e}")
                worksheet.write(1, 1, "N/A")
                worksheet.write(2, 1, "N/A")
            
            # Merge header for sample name
            sample_name = "unknown"
            try:
                if "cluster" in safe_data.columns and len(safe_data["cluster"]) > 0:
                    sample_name = str(safe_data["cluster"].iloc[0]).split("_", 1)[0]
            except Exception:
                sample_name = "unknown"
            worksheet.merge_range(0, 0, 0, 1, sample_name, merge_format)
            
            # Set column widths
            worksheet.set_column(2, 2, 30)  # cluster
            worksheet.set_column(3, 3, 50)  # seq
            worksheet.set_column(4, 4, 15)  # pic_path
            worksheet.set_column(5, 5, 12)  # size, %
            worksheet.set_column(6, 6, 40)  # TAREAN_annotation
            
            # BLAST columns
            if 'best_hit' in safe_data.columns:
                worksheet.set_column(7, 7, 40)  # best_hit
                worksheet.set_column(8, 8, 15)  # best_evalue
                worksheet.set_column(9, 9, 12)  # best_pident
                worksheet.set_column(10, 10, 12)  # coverage
                worksheet.set_column(11, 11, 15)  # coverage_type
                worksheet.set_column(12, 12, 20)  # task_type
        
        logger.info(f"Excel report saved to {output_file}")
        return True
    
    def create_word_report(self, index, data, output_file):
        """Create Word document report with detailed BLAST analysis"""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available, cannot create Word report")
            return False
        
        doc = docx.Document()
        doc.add_heading(f'Repeat Analysis Report - {index}', 0)
        
        # Add summary statistics
        doc.add_heading('Summary Statistics', level=1)
        # Guard against missing size column by substituting zeros.
        safe_data = data.copy()
        if 'size, %' not in safe_data.columns:
            safe_data['size, %'] = 0.0

        total_clusters = len(safe_data)
        total_size = safe_data['size, %'].sum()
        doc.add_paragraph(f'Total clusters: {total_clusters}')
        doc.add_paragraph(f'Total genome coverage: {total_size:.2f}%')
        
        # Count BLAST hits
        if 'best_hit' in data.columns:
            clusters_with_hits = data[data['best_hit'] != 'No significant hits'].shape[0]
            doc.add_paragraph(f'Clusters with BLAST hits: {clusters_with_hits}')
            
            # Coverage statistics
            if 'coverage_type' in data.columns:
                coverage_stats = data['coverage_type'].value_counts()
                doc.add_paragraph('Coverage analysis:')
                for coverage_type, count in coverage_stats.items():
                    if coverage_type != 'none':
                        doc.add_paragraph(f'  {coverage_type}: {count} clusters')
        
        # Add cluster details
        doc.add_heading('Cluster Details', level=1)
        for _, row in safe_data.iterrows():
            # Handle both 'cluster' and 'Cluster' column names
            cluster_col = 'cluster' if 'cluster' in row else 'Cluster'
            doc.add_paragraph(f"Cluster: {row[cluster_col]}")
            doc.add_paragraph(f"Size: {row['size, %']:.2f}%")
            
            # Add TAREAN annotation if available
            if 'TAREAN_annotation' in row and pd.notna(row['TAREAN_annotation']):
                doc.add_paragraph(f"TAREAN annotation: {row['TAREAN_annotation']}")
            
            # Add BLAST analysis if available
            if 'best_hit' in row and pd.notna(row['best_hit']):
                if row['best_hit'] != 'No significant hits':
                    doc.add_paragraph(f"Best BLAST hit: {row['best_hit']}")
                    doc.add_paragraph(f"E-value: {row['best_evalue']:.2e}")
                    doc.add_paragraph(f"Percent identity: {row['best_pident']:.1f}%")
                    
                    if 'coverage' in row and pd.notna(row['coverage']):
                        doc.add_paragraph(f"Coverage: {row['coverage']:.1%} ({row['coverage_type']})")
                    
                    if 'task_type' in row and pd.notna(row['task_type']):
                        doc.add_paragraph(f"Homology type: {row['task_type']}")
                    
                    if 'num_significant' in row and pd.notna(row['num_significant']):
                        doc.add_paragraph(f"Significant matches: {row['num_significant']}")
                else:
                    doc.add_paragraph("No significant BLAST hits found")
            
            doc.add_paragraph("")  # Empty line
        
        doc.save(output_file)
        logger.info(f"Word report saved to {output_file}")
        return True
    
    def run_full_pipeline(self, index, output_dir=None, report_only=False):
        """Run the complete analysis pipeline using the step-based orchestrator."""
        logger.info(f"Starting full pipeline for {index}")

        # Use config output directory if not specified
        if output_dir is None:
            output_dir = self.config.analysis.output_dir

        # Resolve the TAREAN path for this index
        tarean_path, candidates = resolve_tarean_path(self.config_manager, self.config, index)
        if tarean_path is None:
            logger.error(f"Could not find TAREAN data for {index}. Tried paths: {candidates}")
            return False

        # Build orchestration context and step registry
        context = PostTareanContext(
            project_id=getattr(self.config, "project_id", None)
            or getattr(self.config_manager, "project_id", None)
            or "unknown",
            subject=index,
            mode="sample",
            tarean_path=tarean_path,
            output_dir=Path(output_dir),
        )

        steps = build_default_steps()
        orchestrator = PostTareanOrchestrator(analyzer=self, context=context, steps=steps)

        # Determine which logical steps to run, based on analysis configuration.
        enabled_steps = list(getattr(self.config.analysis, "enabled_steps", ["blast", "summary"]))
        if report_only:
            # CLI flag takes precedence over config: skip BLAST entirely.
            enabled_steps = [step for step in enabled_steps if step != "blast"]

        try:
            orchestrator.run(enabled_steps=enabled_steps)
            logger.info(f"Pipeline completed successfully for {index}")
            return True
        except Exception as e:
            logger.error(f"Pipeline failed for {index}: {str(e)}")
            logger.error(f"Traceback: {traceback.format_exc()}")
            return False


def main():
    parser = argparse.ArgumentParser(description='RepOrtR Post-TAREAN Pipeline')
    parser.add_argument('sample', help='Sample name')
    parser.add_argument('--project-id', required=True, help='Project ID')
    parser.add_argument('--output-dir', help='Output directory')
    parser.add_argument('--report-only', action='store_true', help='Generate reports only (skip BLAST)')
    parser.add_argument('--verbose', action='store_true', help='Verbose output')
    parser.add_argument('--test-mode', action='store_true', help='Run in test mode with smaller datasets')
    
    args = parser.parse_args()
    
    # Setup logging
    log_level = logging.INFO if args.verbose else logging.WARNING
    logging.basicConfig(level=log_level, format='%(asctime)s - %(levelname)s - %(message)s')
    
    try:
        # Load configuration (project-scoped)
        config_manager = ConfigManager()
        config = config_manager.load_post_tarean_config(args.project_id)
        
        if args.verbose:
            logger.info(f"Loaded post-TAREAN config for project {args.project_id} from global config")
        
        # Create analyzer using the resolved config
        analyzer = RepeatAnalyzer(config=config, project_id=args.project_id)

        # Adjust for test mode
        if args.test_mode:
            logger.info("Running in test mode with reduced parameters")
            # Reduce memory usage and other parameters for testing
            if hasattr(config, 'performance'):
                config.performance.max_memory_mb = min(config.performance.max_memory_mb, 1024)
        
        # Run pipeline
        ok = analyzer.run_full_pipeline(
            args.sample,
            args.output_dir,
            report_only=bool(args.report_only),
        )
        if not ok:
            logger.error(f"Pipeline reported failure status for {args.sample}")
            sys.exit(1)
        
        logger.info(f"Pipeline completed successfully for {args.sample}")
        
    except Exception as e:
        logger.error(f"Pipeline failed for {args.sample}: {e}")
        if args.verbose:
            logger.error(f"Traceback: {traceback.format_exc()}")
        sys.exit(1)


if __name__ == "__main__":
    main() 